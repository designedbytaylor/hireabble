"""
File uploads routes for Hireabble API — Supabase Storage
Includes media tracking and automatic content moderation.
"""
from fastapi import APIRouter, HTTPException, Depends, UploadFile, File
from fastapi.responses import RedirectResponse
from supabase import create_client
from datetime import datetime, timezone
import uuid
import os
import io
import re

from database import (
    db, get_current_user, SUPABASE_URL, SUPABASE_KEY, UPLOADS_DIR, logger
)

router = APIRouter(tags=["Uploads"])

ALLOWED_IMAGE_TYPES = ["image/jpeg", "image/png", "image/gif", "image/webp"]
ALLOWED_VIDEO_TYPES = ["video/mp4", "video/webm", "video/quicktime", "video/x-msvideo"]
MAX_IMAGE_SIZE = 5 * 1024 * 1024  # 5MB
MAX_VIDEO_SIZE = 50 * 1024 * 1024  # 50MB

# Magic number signatures for file type validation
_IMAGE_MAGIC = {
    b'\xff\xd8\xff': "image/jpeg",
    b'\x89PNG': "image/png",
    b'GIF87a': "image/gif",
    b'GIF89a': "image/gif",
    b'RIFF': "image/webp",  # WebP starts with RIFF...WEBP
}

_VIDEO_MAGIC = {
    b'\x00\x00\x00': "video/mp4",      # ftyp box (MP4/MOV)
    b'\x1a\x45\xdf\xa3': "video/webm", # EBML header (WebM/MKV)
}


def _validate_file_magic(contents: bytes, expected_types: list) -> bool:
    """Validate file contents match expected type via magic numbers."""
    for magic, mime in {**_IMAGE_MAGIC, **_VIDEO_MAGIC}.items():
        if contents[:len(magic)] == magic and mime in expected_types:
            return True
    return False

PHOTO_BUCKET = "photos"
VIDEO_BUCKET = "videos"

# Whether Supabase storage is configured
_USE_SUPABASE = bool(SUPABASE_URL and SUPABASE_KEY)


def _get_supabase():
    if not _USE_SUPABASE:
        raise HTTPException(status_code=500, detail="Storage not configured")
    return create_client(SUPABASE_URL, SUPABASE_KEY)


def _public_url(bucket: str, path: str) -> str:
    """Build the public URL for a stored object."""
    return f"{SUPABASE_URL}/storage/v1/object/public/{bucket}/{path}"


def _save_local(subdir: str, filename: str, contents: bytes) -> str:
    """Save file to local uploads directory, return relative URL path."""
    target_dir = UPLOADS_DIR / subdir
    target_dir.mkdir(parents=True, exist_ok=True)
    filepath = target_dir / filename
    filepath.write_bytes(contents)
    return f"/uploads/{subdir}/{filename}"


def _analyze_image(contents: bytes) -> dict:
    """
    Basic image analysis for content moderation.
    Returns a dict with analysis results and whether to auto-flag.
    Uses PIL to check image properties. Falls back gracefully if PIL unavailable.
    """
    result = {"flagged": False, "reasons": [], "analysis": {}}

    try:
        from PIL import Image
        img = Image.open(io.BytesIO(contents))
        result["analysis"]["width"] = img.width
        result["analysis"]["height"] = img.height
        result["analysis"]["format"] = img.format

        # Flag suspiciously small images (likely spam/placeholder)
        if img.width < 10 or img.height < 10:
            result["flagged"] = True
            result["reasons"].append("Suspiciously small image dimensions")

        # Flag very large images that might be trying to exploit the system
        if img.width > 8000 or img.height > 8000:
            result["flagged"] = True
            result["reasons"].append("Extremely large image dimensions")

        # Check for animated GIFs (often used for inappropriate content)
        if img.format == "GIF":
            try:
                img.seek(1)
                result["analysis"]["animated"] = True
                # Animated GIFs get flagged for review (common vector for inappropriate content)
                result["flagged"] = True
                result["reasons"].append("Animated GIF - requires manual review")
            except EOFError:
                result["analysis"]["animated"] = False

        # Skin-tone pixel ratio analysis (basic NSFW heuristic)
        # This is a lightweight check - not a replacement for ML-based moderation
        if img.mode in ("RGB", "RGBA") and img.format != "GIF":
            try:
                # Sample a subset of pixels for performance
                small = img.resize((100, 100))
                pixels = list(small.getdata())
                skin_count = 0
                for pixel in pixels:
                    r, g, b = pixel[0], pixel[1], pixel[2]
                    # Skin tone detection heuristic
                    if (r > 95 and g > 40 and b > 20 and
                        max(r, g, b) - min(r, g, b) > 15 and
                        abs(r - g) > 15 and r > g and r > b):
                        skin_count += 1
                skin_ratio = skin_count / len(pixels)
                result["analysis"]["skin_ratio"] = round(skin_ratio, 3)
                # High skin-tone ratio flags for review
                if skin_ratio > 0.6:
                    result["flagged"] = True
                    result["reasons"].append(f"High skin-tone ratio ({skin_ratio:.0%}) - flagged for review")
            except Exception:
                pass

    except ImportError:
        # PIL not available - skip analysis, don't flag
        result["analysis"]["note"] = "PIL not installed - image analysis skipped"
    except Exception as e:
        logger.warning(f"Image analysis failed: {e}")
        result["analysis"]["error"] = str(e)

    return result


async def _track_upload(user_id: str, user_name: str, media_type: str, category: str,
                        url: str, filename: str, file_size: int, content_type: str,
                        analysis: dict = None):
    """Track upload in media_uploads collection and auto-flag to moderation if needed."""
    flagged = analysis.get("flagged", False) if analysis else False
    reasons = analysis.get("reasons", []) if analysis else []

    upload_doc = {
        "id": str(uuid.uuid4()),
        "user_id": user_id,
        "user_name": user_name,
        "media_type": media_type,  # 'image' or 'video'
        "category": category,       # 'profile_photo', 'video_intro', 'chat_image', 'chat_video'
        "url": url,
        "filename": filename,
        "file_size": file_size,
        "content_type": content_type,
        "status": "flagged" if flagged else "approved",  # auto-approve clean uploads
        "flagged": flagged,
        "flag_reasons": reasons,
        "analysis": analysis.get("analysis", {}) if analysis else {},
        "created_at": datetime.now(timezone.utc).isoformat(),
        "reviewed_at": None,
        "reviewed_by": None,
    }
    await db.media_uploads.insert_one(upload_doc)

    # If flagged, also add to the moderation queue
    if flagged:
        mod_doc = {
            "id": str(uuid.uuid4()),
            "content_type": "media",
            "content_id": upload_doc["id"],
            "user_id": user_id,
            "status": "pending",
            "violations": [{"category": "image_moderation", "word": r, "severity": "critical"} for r in reasons],
            "created_at": datetime.now(timezone.utc).isoformat(),
            "metadata": {
                "media_url": url,
                "media_category": category,
                "filename": filename,
            }
        }
        await db.moderation_queue.insert_one(mod_doc)
        logger.warning(f"Media flagged for review: {filename} by user {user_id} - {reasons}")

    return upload_doc["id"]


@router.post("/upload/photo")
@router.post("/upload")
async def upload_file(
    file: UploadFile = File(...),
    current_user: dict = Depends(get_current_user)
):
    """Upload a profile photo (Supabase Storage or local fallback)"""
    if file.content_type not in ALLOWED_IMAGE_TYPES:
        raise HTTPException(status_code=400, detail="Only image files are allowed")

    contents = await file.read()
    if len(contents) > MAX_IMAGE_SIZE:
        raise HTTPException(status_code=400, detail="File size exceeds 5MB limit")

    if not _validate_file_magic(contents, ALLOWED_IMAGE_TYPES):
        raise HTTPException(status_code=400, detail="File content does not match an allowed image type")

    # Analyze image content
    analysis = _analyze_image(contents)

    ext = file.filename.split('.')[-1] if '.' in file.filename else 'jpg'
    filename = f"{current_user['id']}_{uuid.uuid4().hex[:8]}.{ext}"

    if _USE_SUPABASE:
        supabase = _get_supabase()

        # Remove old photo if it exists in the bucket
        user = await db.users.find_one({"id": current_user["id"]})
        if user and user.get("photo_url") and SUPABASE_URL in (user.get("photo_url") or ""):
            old_path = user["photo_url"].split(f"/{PHOTO_BUCKET}/")[-1]
            try:
                supabase.storage.from_(PHOTO_BUCKET).remove([old_path])
            except Exception:
                pass

        try:
            supabase.storage.from_(PHOTO_BUCKET).upload(
                filename,
                contents,
                file_options={"content-type": file.content_type}
            )
        except Exception as e:
            logger.error(f"Supabase upload failed: {e}")
            raise HTTPException(status_code=500, detail="Failed to upload photo")

        photo_url = _public_url(PHOTO_BUCKET, filename)
    else:
        # Local fallback
        photo_url = _save_local("photos", filename, contents)

    await db.users.update_one(
        {"id": current_user["id"]},
        {"$set": {"photo_url": photo_url}}
    )

    # Track upload
    await _track_upload(
        user_id=current_user["id"],
        user_name=current_user.get("name", "Unknown"),
        media_type="image",
        category="profile_photo",
        url=photo_url,
        filename=filename,
        file_size=len(contents),
        content_type=file.content_type,
        analysis=analysis
    )

    logger.info(f"Photo uploaded for user {current_user['id']}: {filename}")
    return {"photo_url": photo_url, "filename": filename}


@router.post("/upload/video")
async def upload_video(
    file: UploadFile = File(...),
    current_user: dict = Depends(get_current_user)
):
    """Upload a video introduction to Supabase Storage"""
    if current_user["role"] != "seeker":
        raise HTTPException(status_code=403, detail="Only job seekers can upload video intros")

    if file.content_type not in ALLOWED_VIDEO_TYPES:
        raise HTTPException(status_code=400, detail="Only video files (MP4, WebM, MOV) are allowed")

    contents = await file.read()
    if len(contents) > MAX_VIDEO_SIZE:
        raise HTTPException(status_code=400, detail="Video size exceeds 50MB limit")

    if not _validate_file_magic(contents, ALLOWED_VIDEO_TYPES):
        raise HTTPException(status_code=400, detail="File content does not match an allowed video type")

    ext = file.filename.split('.')[-1] if '.' in file.filename else 'mp4'
    filename = f"video_{current_user['id']}_{uuid.uuid4().hex[:8]}.{ext}"

    if _USE_SUPABASE:
        supabase = _get_supabase()

        user = await db.users.find_one({"id": current_user["id"]})
        if user and user.get("video_url") and SUPABASE_URL in (user.get("video_url") or ""):
            old_path = user["video_url"].split(f"/{VIDEO_BUCKET}/")[-1]
            try:
                supabase.storage.from_(VIDEO_BUCKET).remove([old_path])
            except Exception:
                pass

        try:
            supabase.storage.from_(VIDEO_BUCKET).upload(
                filename,
                contents,
                file_options={"content-type": file.content_type}
            )
        except Exception as e:
            logger.error(f"Supabase video upload failed: {e}")
            raise HTTPException(status_code=500, detail="Failed to upload video")

        video_url = _public_url(VIDEO_BUCKET, filename)
    else:
        video_url = _save_local("videos", filename, contents)

    await db.users.update_one(
        {"id": current_user["id"]},
        {"$set": {"video_url": video_url}}
    )

    # Track upload (no image analysis for video — flagged for manual review by default if needed)
    await _track_upload(
        user_id=current_user["id"],
        user_name=current_user.get("name", "Unknown"),
        media_type="video",
        category="video_intro",
        url=video_url,
        filename=filename,
        file_size=len(contents),
        content_type=file.content_type,
    )

    logger.info(f"Video uploaded for user {current_user['id']}: {filename}")
    return {"video_url": video_url, "filename": filename}


@router.delete("/upload/video")
async def delete_video(current_user: dict = Depends(get_current_user)):
    """Delete video introduction from Supabase Storage"""
    if current_user["role"] != "seeker":
        raise HTTPException(status_code=403, detail="Only job seekers can manage video intros")

    user = await db.users.find_one({"id": current_user["id"]})
    if user and user.get("video_url"):
        # Delete from Supabase if it's a Supabase URL
        if SUPABASE_URL and SUPABASE_URL in (user.get("video_url") or ""):
            old_path = user["video_url"].split(f"/{VIDEO_BUCKET}/")[-1]
            try:
                supabase = _get_supabase()
                supabase.storage.from_(VIDEO_BUCKET).remove([old_path])
            except Exception:
                pass

        await db.users.update_one(
            {"id": current_user["id"]},
            {"$set": {"video_url": None}}
        )

        # Mark in media_uploads as removed
        await db.media_uploads.update_many(
            {"user_id": current_user["id"], "category": "video_intro"},
            {"$set": {"status": "removed", "reviewed_at": datetime.now(timezone.utc).isoformat()}}
        )

        logger.info(f"Video deleted for user {current_user['id']}")

    return {"message": "Video deleted successfully"}


@router.post("/upload/chat-image")
async def upload_chat_image(
    file: UploadFile = File(...),
    current_user: dict = Depends(get_current_user)
):
    """Upload an image for chat messages"""
    if file.content_type not in ALLOWED_IMAGE_TYPES:
        raise HTTPException(status_code=400, detail="Only image files are allowed")

    contents = await file.read()
    if len(contents) > MAX_IMAGE_SIZE:
        raise HTTPException(status_code=400, detail="Image size exceeds 5MB limit")

    # Analyze image content
    analysis = _analyze_image(contents)

    ext = file.filename.split('.')[-1] if '.' in file.filename else 'jpg'
    filename = f"chat_{current_user['id']}_{uuid.uuid4().hex[:8]}.{ext}"

    if _USE_SUPABASE:
        supabase = _get_supabase()
        try:
            supabase.storage.from_(PHOTO_BUCKET).upload(
                filename,
                contents,
                file_options={"content-type": file.content_type}
            )
        except Exception as e:
            logger.error(f"Supabase chat image upload failed: {e}")
            raise HTTPException(status_code=500, detail="Failed to upload image")
        image_url = _public_url(PHOTO_BUCKET, filename)
    else:
        image_url = _save_local("photos", filename, contents)

    # Track upload
    await _track_upload(
        user_id=current_user["id"],
        user_name=current_user.get("name", "Unknown"),
        media_type="image",
        category="chat_image",
        url=image_url,
        filename=filename,
        file_size=len(contents),
        content_type=file.content_type,
        analysis=analysis
    )

    logger.info(f"Chat image uploaded by user {current_user['id']}: {filename}")
    return {"url": image_url, "filename": filename}


@router.post("/upload/chat-video")
async def upload_chat_video(
    file: UploadFile = File(...),
    current_user: dict = Depends(get_current_user)
):
    """Upload a video message for chat"""
    if file.content_type not in ALLOWED_VIDEO_TYPES:
        raise HTTPException(status_code=400, detail="Only video files (MP4, WebM, MOV) are allowed")

    contents = await file.read()
    if len(contents) > MAX_VIDEO_SIZE:
        raise HTTPException(status_code=400, detail="Video size exceeds 50MB limit")

    ext = file.filename.split('.')[-1] if '.' in file.filename else 'webm'
    filename = f"chat_{current_user['id']}_{uuid.uuid4().hex[:8]}.{ext}"

    if _USE_SUPABASE:
        supabase = _get_supabase()
        try:
            supabase.storage.from_(VIDEO_BUCKET).upload(
                filename,
                contents,
                file_options={"content-type": file.content_type}
            )
        except Exception as e:
            logger.error(f"Supabase chat video upload failed: {e}")
            raise HTTPException(status_code=500, detail="Failed to upload video")
        video_url = _public_url(VIDEO_BUCKET, filename)
    else:
        video_url = _save_local("videos", filename, contents)

    # Track upload
    await _track_upload(
        user_id=current_user["id"],
        user_name=current_user.get("name", "Unknown"),
        media_type="video",
        category="chat_video",
        url=video_url,
        filename=filename,
        file_size=len(contents),
        content_type=file.content_type,
    )

    logger.info(f"Chat video uploaded by user {current_user['id']}: {filename}")
    return {"url": video_url, "filename": filename}


# ==================== RESUME PARSING ====================

MAX_RESUME_SIZE = 10 * 1024 * 1024  # 10MB
ALLOWED_RESUME_TYPES = [
    "application/pdf",
    "application/msword",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
]

async def _parse_resume_with_ai(text: str) -> dict:
    """Extract structured data from resume text using Claude AI for accurate parsing."""
    import json as json_module

    default_result = {
        "name": None,
        "title": None,
        "email": None,
        "phone": None,
        "location": None,
        "skills": [],
        "work_history": [],
        "education": [],
        "certifications": [],
        "bio": None,
    }

    if not text.strip():
        return default_result

    api_key = os.environ.get("ANTHROPIC_API_KEY")
    if not api_key:
        logger.warning("ANTHROPIC_API_KEY not set — falling back to basic parsing")
        result = _parse_resume_basic(text)
        result["_parser"] = "basic"
        result["_parser_reason"] = "ANTHROPIC_API_KEY not set"
        return result

    try:
        import anthropic
    except ImportError:
        logger.error("anthropic package not installed — falling back to basic parsing. Install with: pip install anthropic")
        result = _parse_resume_basic(text)
        result["_parser"] = "basic"
        result["_parser_reason"] = "anthropic package not installed"
        return result

    try:
        client = anthropic.Anthropic(api_key=api_key)

        prompt = f"""You are a resume parser. Extract structured data from the resume text below. Return ONLY valid JSON, no extra text.

Required JSON structure:
{{
  "name": "Full Name",
  "title": "Most recent or primary job title",
  "email": "email@example.com",
  "phone": "phone number",
  "location": "City, State",
  "skills": ["skill1", "skill2"],
  "work_history": [
    {{
      "company": "Company Name",
      "position": "Job Title",
      "start_date": "Jan 2020",
      "end_date": "Present",
      "description": "Combined bullet points describing responsibilities and achievements"
    }}
  ],
  "education": [
    {{
      "school": "University Name",
      "degree": "Bachelor of Science",
      "field": "Computer Science",
      "year": "2020"
    }}
  ],
  "certifications": ["cert1", "cert2"],
  "bio": "Brief professional summary if one exists, otherwise null",
  "experience_years": 5
}}

CRITICAL PARSING RULES:

WORK HISTORY — This is the most important section:
- Every job listed on the resume MUST be a separate entry in work_history
- "position" = the job title (e.g., "Software Engineer", "Course Coordinator", "IT Systems Intern")
- "company" = the employer name (e.g., "Google", "NP Photonics Inc", "Statefarm Insurance")
- "start_date" = formatted as "Mon YYYY" (e.g., "May 2017", "Aug 2018"). If only a year, use "Jan YYYY"
- "end_date" = formatted as "Mon YYYY" or "Present" if current/ongoing
- "description" = Combine ALL bullet points for that position into a single string, separated by newlines. Include the actual content of each bullet point. Do NOT summarize — preserve the original descriptions.
- Order: most recent job first
- Research positions, internships, teaching roles, and volunteer work all count as work history entries

EDUCATION:
- Each school/degree is a separate entry
- "degree" = the degree type spelled out (e.g., "Bachelor of Science", "Master of Arts", "Associate of Science")
- "field" = the major/field of study (e.g., "Computer Science", "Mathematics")
- "year" = graduation year as string, or "Expected YYYY" if not yet graduated
- Include GPA in the field if mentioned (e.g., "Computer Science (GPA: 3.25/4.00)")
- Awards like "Dean's List" should be appended to the field

SKILLS:
- Extract from the skills/technologies section
- Include programming languages, frameworks, tools, methodologies, soft skills
- Max 30 skills, most relevant first
- Do NOT extract random words from job descriptions

TITLE: Use the most recent job title, NOT the person's name.

EXPERIENCE_YEARS: Calculate total years of professional experience from work history dates. Round to nearest integer.

CERTIFICATIONS: Include professional certifications, licenses, and significant awards/honors.

Use null for fields you cannot confidently determine.

Resume text:
{text}"""

        # Try models in order of preference
        models = [
            "claude-haiku-4-5-20251001",
            "claude-3-5-haiku-20241022",
            "claude-3-haiku-20240307",
        ]
        message = None
        last_error = None
        for model_id in models:
            try:
                message = client.messages.create(
                    model=model_id,
                    max_tokens=4000,
                    messages=[{"role": "user", "content": prompt}],
                )
                logger.info(f"AI resume parsing used model: {model_id}")
                break
            except Exception as model_err:
                last_error = model_err
                logger.warning(f"Model {model_id} failed: {type(model_err).__name__}: {model_err}")
                continue

        if message is None:
            raise last_error or Exception("All models failed")

        response_text = message.content[0].text.strip()

        # Extract JSON from response (handle potential markdown code blocks)
        if response_text.startswith("```"):
            # Remove markdown code block wrapper
            lines = response_text.split("\n")
            json_lines = []
            in_block = False
            for line in lines:
                if line.startswith("```") and not in_block:
                    in_block = True
                    continue
                elif line.startswith("```") and in_block:
                    break
                elif in_block:
                    json_lines.append(line)
            response_text = "\n".join(json_lines)

        parsed = json_module.loads(response_text)

        # Validate and ensure correct structure
        result = {
            "name": parsed.get("name"),
            "title": parsed.get("title"),
            "email": parsed.get("email"),
            "phone": parsed.get("phone"),
            "location": parsed.get("location"),
            "skills": parsed.get("skills", [])[:30],
            "work_history": [],
            "education": [],
            "certifications": parsed.get("certifications", [])[:15],
            "bio": parsed.get("bio"),
            "experience_years": parsed.get("experience_years"),
        }

        # Validate work_history entries
        for entry in parsed.get("work_history", [])[:10]:
            if isinstance(entry, dict):
                result["work_history"].append({
                    "company": entry.get("company", ""),
                    "position": entry.get("position", ""),
                    "start_date": entry.get("start_date", ""),
                    "end_date": entry.get("end_date", ""),
                    "description": entry.get("description", ""),
                })

        # Validate education entries
        for entry in parsed.get("education", [])[:5]:
            if isinstance(entry, dict):
                result["education"].append({
                    "school": entry.get("school", ""),
                    "degree": entry.get("degree", ""),
                    "field": entry.get("field", ""),
                    "year": entry.get("year", ""),
                })

        logger.info(f"AI resume parsing successful: {len(result['skills'])} skills, {len(result['work_history'])} positions, {len(result['education'])} education")
        result["_parser"] = "ai"
        return result

    except Exception as e:
        logger.error(f"AI resume parsing failed: {type(e).__name__}: {e} — falling back to basic parsing")
        result = _parse_resume_basic(text)
        result["_parser"] = "basic"
        result["_parser_reason"] = f"{type(e).__name__}: {str(e)[:200]}"
        return result


def _parse_resume_basic(text: str) -> dict:
    """Basic fallback resume parsing using section-based pattern matching when AI is unavailable."""
    result = {
        "name": None,
        "title": None,
        "email": None,
        "phone": None,
        "location": None,
        "skills": [],
        "work_history": [],
        "education": [],
        "certifications": [],
        "bio": None,
    }

    lines = [line.strip() for line in text.split('\n') if line.strip()]
    if not lines:
        return result

    # Extract email
    email_match = re.search(r'[\w.+-]+@[\w-]+\.[\w.-]+', text)
    if email_match:
        result["email"] = email_match.group(0)

    # Extract phone
    phone_match = re.search(r'[\+]?[\d\s\-\(\)]{10,15}', text)
    if phone_match:
        candidate = phone_match.group(0).strip()
        digits = re.sub(r'\D', '', candidate)
        if 10 <= len(digits) <= 15:
            result["phone"] = candidate

    # Name: first non-empty line that's not email/phone/header keyword
    header_words = {'experience', 'education', 'skills', 'summary', 'objective', 'profile', 'certifications'}
    name_candidates = [l for l in lines[:5]
                       if not re.search(r'@|http|www\.|[\d\s\-\(\)]{10,}', l)
                       and len(l) < 60
                       and l.lower().strip(':') not in header_words
                       and not re.search(r'bachelor|master|university|college|school|degree', l, re.IGNORECASE)]
    if name_candidates:
        result["name"] = name_candidates[0]

    # Location: "City, ST" pattern — must NOT contain degree/education words
    for loc_match in re.finditer(
        r'([A-Z][a-z]+(?:\s[A-Z][a-z]+)*,\s*[A-Z]{2})(?:\s+\d{5})?',
        text[:600]
    ):
        candidate_loc = loc_match.group(1)
        # Skip if it's inside a degree/education phrase
        start = max(0, loc_match.start() - 40)
        context = text[start:loc_match.end() + 10]
        if not re.search(r'bachelor|master|degree|science|arts|university|college|gpa', context, re.IGNORECASE):
            result["location"] = candidate_loc
            break

    # --- Section-based parsing ---
    section_headers = re.compile(
        r'^(work\s*experience|professional\s*experience|experience|employment\s*(?:history)?|'
        r'education|academic\s*background|skills|technical\s*skills|core\s*competencies|relevant\s*skills|'
        r'certifications?|licenses?\s*(?:and|&)?\s*certifications?|projects?|'
        r'activities|summary|professional\s*summary|objective|'
        r'profile|qualifications|honors?\s*(?:and|&)?\s*awards?)\s*:?\s*$',
        re.IGNORECASE | re.MULTILINE
    )

    # Build sections dict
    sections = {}
    current_section = "header"
    sections[current_section] = []
    for line in lines:
        clean = line.strip()
        if section_headers.match(clean):
            current_section = clean.rstrip(':').strip().lower()
            sections[current_section] = []
        else:
            sections.setdefault(current_section, []).append(clean)

    # --- Parse Work Experience ---
    work_keys = [k for k in sections if any(w in k for w in ['work', 'experience', 'employment'])
                 and 'education' not in k]
    for key in work_keys:
        work_lines = sections[key]
        current_job = None
        pending_title_line = None  # Line before date that might be position/company

        for idx, wline in enumerate(work_lines):
            # Detect a new job entry: line with a date range
            date_match = re.search(
                r'((?:(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)\s+)?\d{4})\s*[-–—]+\s*((?:(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)\s+)?\d{0,4}|[Pp]resent|[Cc]urrent)',
                wline
            )
            if date_match:
                if current_job:
                    current_job["description"] = "\n".join(current_job.get("_bullets", []))
                    current_job.pop("_bullets", None)
                    result["work_history"].append(current_job)

                # Try to extract position and company from the line (text before the date)
                text_before = wline[:date_match.start()].strip().rstrip(',').rstrip('–').rstrip('-').rstrip('|').strip()
                text_after = wline[date_match.end():].strip().lstrip(',').lstrip('|').strip()

                # Split by multiple spaces or tabs
                parts = [p.strip() for p in re.split(r'\s{2,}|\t', text_before) if p.strip()] if text_before else []
                position = ""
                company = ""

                if len(parts) >= 2:
                    position = parts[0]
                    company = parts[1]
                elif len(parts) == 1:
                    # Single part — check previous line for position/company
                    position = parts[0]
                    if pending_title_line:
                        # Previous line might have the other half
                        prev_parts = [p.strip() for p in re.split(r'\s{2,}|\t|[,|]', pending_title_line) if p.strip()]
                        if prev_parts:
                            # If current looks like a company, swap
                            if any(w in position.lower() for w in ['inc', 'corp', 'llc', 'ltd', 'company', 'insurance', 'university']):
                                company = position
                                position = prev_parts[0]
                            else:
                                company = prev_parts[0] if len(prev_parts) == 1 else prev_parts[-1]
                elif pending_title_line:
                    # No text before date — position/company from previous line
                    prev_parts = [p.strip() for p in re.split(r'\s{2,}|\t', pending_title_line) if p.strip()]
                    if len(prev_parts) >= 2:
                        position = prev_parts[0]
                        company = prev_parts[1]
                    elif prev_parts:
                        position = prev_parts[0]

                # Clean bullet markers from position
                position = re.sub(r'^[●•·▪◦○]\s*', '', position)

                current_job = {
                    "position": position,
                    "company": company,
                    "start_date": date_match.group(1).strip(),
                    "end_date": date_match.group(2).strip() if date_match.group(2) else "Present",
                    "_bullets": [],
                }
                pending_title_line = None
            elif current_job:
                # Bullet point or continuation line
                bullet = re.sub(r'^[●•·▪◦○o\-]\s*', '', wline).strip()
                if bullet and len(bullet) > 5:
                    current_job["_bullets"].append(bullet)
            else:
                # Line before first date — could be position/company header
                pending_title_line = wline

        if current_job:
            current_job["description"] = "\n".join(current_job.get("_bullets", []))
            current_job.pop("_bullets", None)
            result["work_history"].append(current_job)

    # If we found work history, use the first position as title
    if result["work_history"] and not result["title"]:
        result["title"] = result["work_history"][0].get("position")

    # Calculate experience years from work history dates
    if result["work_history"]:
        try:
            from datetime import datetime as dt
            earliest_year = None
            for job in result["work_history"]:
                year_match = re.search(r'\d{4}', job.get("start_date", ""))
                if year_match:
                    y = int(year_match.group())
                    if earliest_year is None or y < earliest_year:
                        earliest_year = y
            if earliest_year:
                result["experience_years"] = max(1, dt.now().year - earliest_year)
        except Exception:
            pass

    # --- Parse Education ---
    edu_keys = [k for k in sections if any(w in k for w in ['education', 'academic'])]
    for key in edu_keys:
        edu_lines = sections[key]
        degree_pattern = re.compile(
            r'(bachelor|master|associate|ph\.?d|doctor|b\.?s\.?|b\.?a\.?|m\.?s\.?|m\.?a\.?|m\.?b\.?a)',
            re.IGNORECASE
        )
        i = 0
        while i < len(edu_lines):
            line = edu_lines[i]
            has_degree = degree_pattern.search(line)
            has_school = any(w in line.lower() for w in ['university', 'college', 'institute', 'school', 'academy'])

            if has_degree or has_school:
                edu_entry = {"school": "", "degree": "", "field": "", "year": ""}
                # Gather context lines for this education entry
                context_lines = [line]
                j = i + 1
                while j < len(edu_lines) and j < i + 5:
                    next_line = edu_lines[j]
                    # Stop if another school/degree
                    if (any(w in next_line.lower() for w in ['university', 'college', 'institute'])
                            and degree_pattern.search(next_line) and j > i + 1):
                        break
                    context_lines.append(next_line)
                    j += 1

                # Process each context line separately for better field extraction
                for ctx_line in context_lines:
                    # Extract school name
                    if not edu_entry["school"]:
                        school_match = re.search(
                            r'((?:The\s+)?(?:University|College|Institute|School|Academy)\s+of\s+[\w\s]+?(?=\s{2}|\t|\d|,\s*[A-Z]{2}|$)|'
                            r'[\w\s]+?(?:University|College|Institute|School|Academy|Tech))',
                            ctx_line, re.IGNORECASE
                        )
                        if school_match:
                            school_name = school_match.group(0).strip()
                            # Clean trailing whitespace artifacts
                            school_name = re.sub(r'\s+$', '', school_name)
                            edu_entry["school"] = school_name

                    # Extract degree type
                    if not edu_entry["degree"]:
                        deg_match = re.search(
                            r"((?:Bachelor|Master|Associate|Doctor(?:ate)?)\s+of\s+(?:Science|Arts|Engineering|Business|Fine Arts|Education|Applied Science))",
                            ctx_line, re.IGNORECASE
                        )
                        if not deg_match:
                            deg_match = re.search(
                                r"(B\.?S\.?|B\.?A\.?|M\.?S\.?|M\.?A\.?|M\.?B\.?A\.?|Ph\.?D\.?|A\.?S\.?)",
                                ctx_line
                            )
                        if deg_match:
                            edu_entry["degree"] = deg_match.group(0).strip()

                    # Extract field of study — look for "in <field>" pattern
                    if not edu_entry["field"]:
                        field_match = re.search(
                            r'(?:Bachelor|Master|B\.?S\.?|B\.?A\.?|M\.?S\.?|M\.?A\.?|Ph\.?D\.?)\s+(?:of\s+\w+\s+)?in\s+([\w\s&]+?)(?:\s*[\(,|]|\s{2,}|\s*$)',
                            ctx_line, re.IGNORECASE
                        )
                        if field_match:
                            edu_entry["field"] = field_match.group(1).strip()[:60]
                        else:
                            # Try "Major: X" or standalone field after degree
                            major_match = re.search(r'[Mm]ajor:?\s*([\w\s&]+?)(?:\s*[,;|]|\s*$)', ctx_line)
                            if major_match:
                                edu_entry["field"] = major_match.group(1).strip()[:60]

                    # Extract year
                    if not edu_entry["year"]:
                        yr_match = re.search(r'(?:Expected\s+|Graduation:?\s*)?(?:(?:May|June?|Aug(?:ust)?|Dec(?:ember)?|Jan(?:uary)?|Spring|Fall|Summer)\s+)?(\d{4})', ctx_line, re.IGNORECASE)
                        if yr_match:
                            edu_entry["year"] = yr_match.group(0).strip()

                # Extract GPA and add to field (not degree)
                full_context = " ".join(context_lines)
                gpa_match = re.search(r'GPA:?\s*([\d.]+(?:\s*/\s*[\d.]+)?)', full_context, re.IGNORECASE)
                if gpa_match:
                    if edu_entry["field"]:
                        edu_entry["field"] += f" (GPA: {gpa_match.group(1)})"
                    elif edu_entry["degree"]:
                        edu_entry["degree"] += f" (GPA: {gpa_match.group(1)})"

                if edu_entry["school"] or edu_entry["degree"]:
                    result["education"].append(edu_entry)
                i = j
            else:
                i += 1

    # --- Parse Skills ---
    # Known technology/skill keywords for matching across entire resume
    KNOWN_SKILLS = {
        # Programming languages
        'python', 'java', 'javascript', 'typescript', 'c++', 'c#', 'ruby', 'go', 'golang',
        'rust', 'swift', 'kotlin', 'scala', 'php', 'perl', 'r', 'matlab', 'dart', 'lua',
        'haskell', 'elixir', 'clojure', 'objective-c', 'shell', 'bash', 'powershell',
        'sql', 'html', 'css', 'sass', 'less', 'graphql', 'solidity',
        # Frontend frameworks
        'react', 'react.js', 'reactjs', 'angular', 'angularjs', 'vue', 'vue.js', 'vuejs',
        'svelte', 'next.js', 'nextjs', 'nuxt', 'nuxt.js', 'gatsby', 'ember', 'backbone',
        'jquery', 'bootstrap', 'tailwind', 'tailwindcss', 'material-ui', 'mui',
        'chakra ui', 'redux', 'mobx', 'webpack', 'vite', 'rollup', 'parcel',
        # Backend frameworks
        'node.js', 'nodejs', 'express', 'express.js', 'django', 'flask', 'fastapi',
        'spring', 'spring boot', 'rails', 'ruby on rails', 'laravel', 'asp.net', '.net',
        'dotnet', 'gin', 'fiber', 'actix', 'nest.js', 'nestjs', 'koa', 'hapi',
        # Databases
        'mysql', 'postgresql', 'postgres', 'mongodb', 'redis', 'elasticsearch',
        'sqlite', 'oracle', 'sql server', 'dynamodb', 'cassandra', 'neo4j',
        'mariadb', 'couchdb', 'firestore', 'supabase', 'prisma',
        # Cloud & DevOps
        'aws', 'amazon web services', 'azure', 'gcp', 'google cloud', 'heroku',
        'digitalocean', 'vercel', 'netlify', 'docker', 'kubernetes', 'k8s',
        'terraform', 'ansible', 'jenkins', 'circleci', 'github actions',
        'gitlab ci', 'travis ci', 'nginx', 'apache', 'linux', 'unix',
        # Data & ML
        'pandas', 'numpy', 'scikit-learn', 'tensorflow', 'pytorch', 'keras',
        'spark', 'hadoop', 'airflow', 'kafka', 'rabbitmq', 'celery',
        'tableau', 'power bi', 'looker', 'dbt', 'snowflake', 'bigquery',
        'machine learning', 'deep learning', 'nlp', 'computer vision',
        'data analysis', 'data engineering', 'data science', 'etl',
        # Mobile
        'react native', 'flutter', 'ios', 'android', 'swiftui', 'xamarin',
        'ionic', 'cordova', 'expo',
        # Tools & Practices
        'git', 'github', 'gitlab', 'bitbucket', 'jira', 'confluence',
        'figma', 'sketch', 'adobe xd', 'postman', 'swagger',
        'rest', 'restful', 'rest api', 'graphql', 'grpc', 'websocket',
        'ci/cd', 'agile', 'scrum', 'kanban', 'tdd', 'bdd',
        'microservices', 'serverless', 'oauth', 'jwt',
        # Testing
        'jest', 'mocha', 'chai', 'cypress', 'selenium', 'playwright',
        'pytest', 'junit', 'rspec', 'testing library',
        # Other
        'api', 'seo', 'accessibility', 'a11y', 'responsive design',
        'ux', 'ui', 'ux design', 'ui design', 'wireframing',
        'project management', 'team leadership', 'mentoring',
        'communication', 'problem solving', 'analytical',
        'salesforce', 'hubspot', 'sap', 'erp', 'crm',
    }

    # Build a case-insensitive lookup (maps lowercase -> preferred casing)
    SKILL_LOOKUP = {}
    for sk in KNOWN_SKILLS:
        SKILL_LOOKUP[sk.lower()] = sk

    # Canonical casing for well-known skills
    SKILL_DISPLAY = {
        'python': 'Python', 'java': 'Java', 'javascript': 'JavaScript',
        'typescript': 'TypeScript', 'c++': 'C++', 'c#': 'C#', 'ruby': 'Ruby',
        'go': 'Go', 'golang': 'Go', 'rust': 'Rust', 'swift': 'Swift',
        'kotlin': 'Kotlin', 'scala': 'Scala', 'php': 'PHP', 'perl': 'Perl',
        'r': 'R', 'matlab': 'MATLAB', 'dart': 'Dart', 'sql': 'SQL',
        'html': 'HTML', 'css': 'CSS', 'sass': 'Sass', 'graphql': 'GraphQL',
        'react': 'React', 'react.js': 'React', 'reactjs': 'React',
        'angular': 'Angular', 'angularjs': 'Angular',
        'vue': 'Vue.js', 'vue.js': 'Vue.js', 'vuejs': 'Vue.js',
        'svelte': 'Svelte', 'next.js': 'Next.js', 'nextjs': 'Next.js',
        'nuxt': 'Nuxt.js', 'nuxt.js': 'Nuxt.js',
        'gatsby': 'Gatsby', 'jquery': 'jQuery', 'bootstrap': 'Bootstrap',
        'tailwind': 'Tailwind CSS', 'tailwindcss': 'Tailwind CSS',
        'redux': 'Redux', 'webpack': 'Webpack', 'vite': 'Vite',
        'node.js': 'Node.js', 'nodejs': 'Node.js',
        'express': 'Express.js', 'express.js': 'Express.js',
        'django': 'Django', 'flask': 'Flask', 'fastapi': 'FastAPI',
        'spring': 'Spring', 'spring boot': 'Spring Boot',
        'rails': 'Ruby on Rails', 'ruby on rails': 'Ruby on Rails',
        'laravel': 'Laravel', 'asp.net': 'ASP.NET', '.net': '.NET', 'dotnet': '.NET',
        'nest.js': 'NestJS', 'nestjs': 'NestJS',
        'mysql': 'MySQL', 'postgresql': 'PostgreSQL', 'postgres': 'PostgreSQL',
        'mongodb': 'MongoDB', 'redis': 'Redis', 'elasticsearch': 'Elasticsearch',
        'sqlite': 'SQLite', 'dynamodb': 'DynamoDB', 'cassandra': 'Cassandra',
        'neo4j': 'Neo4j', 'firestore': 'Firestore', 'supabase': 'Supabase',
        'prisma': 'Prisma', 'snowflake': 'Snowflake', 'bigquery': 'BigQuery',
        'aws': 'AWS', 'azure': 'Azure', 'gcp': 'GCP',
        'heroku': 'Heroku', 'vercel': 'Vercel', 'docker': 'Docker',
        'kubernetes': 'Kubernetes', 'k8s': 'Kubernetes',
        'terraform': 'Terraform', 'ansible': 'Ansible', 'jenkins': 'Jenkins',
        'nginx': 'Nginx', 'linux': 'Linux',
        'pandas': 'Pandas', 'numpy': 'NumPy', 'tensorflow': 'TensorFlow',
        'pytorch': 'PyTorch', 'keras': 'Keras', 'spark': 'Spark',
        'kafka': 'Kafka', 'tableau': 'Tableau', 'dbt': 'dbt',
        'react native': 'React Native', 'flutter': 'Flutter',
        'ios': 'iOS', 'android': 'Android', 'swiftui': 'SwiftUI',
        'git': 'Git', 'github': 'GitHub', 'gitlab': 'GitLab',
        'jira': 'Jira', 'figma': 'Figma', 'postman': 'Postman',
        'jest': 'Jest', 'cypress': 'Cypress', 'selenium': 'Selenium',
        'playwright': 'Playwright', 'pytest': 'pytest',
        'agile': 'Agile', 'scrum': 'Scrum',
        'salesforce': 'Salesforce', 'sap': 'SAP',
        'ci/cd': 'CI/CD', 'rest': 'REST', 'restful': 'REST',
        'rest api': 'REST API', 'grpc': 'gRPC',
        'machine learning': 'Machine Learning', 'deep learning': 'Deep Learning',
        'nlp': 'NLP', 'computer vision': 'Computer Vision',
        'data analysis': 'Data Analysis', 'data engineering': 'Data Engineering',
        'data science': 'Data Science', 'etl': 'ETL',
        'microservices': 'Microservices', 'serverless': 'Serverless',
    }

    skills_found = set()  # Track lowercase for dedup

    def _add_skill(skill_text):
        """Add a skill with deduplication and normalization."""
        key = skill_text.lower().strip()
        if key in SKILL_DISPLAY:
            display = SKILL_DISPLAY[key]
        else:
            display = skill_text.strip()
        dedup_key = display.lower()
        if dedup_key not in skills_found and len(display) >= 2:
            skills_found.add(dedup_key)
            result["skills"].append(display)

    # 1) Parse dedicated skill sections (original approach, improved header matching)
    skill_keys = [k for k in sections if any(w in k for w in
                  ['skill', 'qualifications', 'competenc', 'technologies', 'proficienc', 'tools', 'expertise'])]
    for key in skill_keys:
        for sline in sections[key]:
            # Remove common prefixes like "Languages:", "Frameworks:", etc.
            clean = re.sub(
                r'^(?:Proficient\s+(?:in|with)|Familiar\s+with|Experience\s+with|'
                r'Languages|Frameworks|Tools|Technologies|Databases|Platforms|'
                r'Libraries|Operating\s+Systems|Software|Hardware|'
                r'Frontend|Backend|DevOps|Cloud|Mobile|Other)\s*[:()]\s*',
                '', sline, flags=re.IGNORECASE
            )
            raw_skills = re.split(r'[,;|●•·▪◦○]', clean)
            for s in raw_skills:
                s = s.strip().strip('()')
                s = re.sub(r'^[●•·▪◦○\-]\s*', '', s).strip()
                if s and 2 <= len(s) <= 40 and not re.match(r'^[\d\s]+$', s):
                    _add_skill(s)

    # 2) Scan work history bullets for inline "Technologies:" lines and known skill mentions
    for job in result["work_history"]:
        desc = job.get("description", "")
        if not desc:
            continue
        for bullet_line in desc.split('\n'):
            # Check for inline tech lists: "Technologies: X, Y, Z" or "Tech stack: ..."
            tech_match = re.match(
                r'(?:Technologies|Tech(?:nology)?\s*(?:Stack|Used)?|Stack|Built\s+with|'
                r'Tools?\s+Used|Environment)\s*[:]\s*(.+)',
                bullet_line, re.IGNORECASE
            )
            if tech_match:
                for s in re.split(r'[,;|●•·]', tech_match.group(1)):
                    s = s.strip().strip('()')
                    if s and 2 <= len(s) <= 40:
                        _add_skill(s)

            # Scan for known skill keywords in bullet text
            bullet_lower = bullet_line.lower()
            for skill_key, display in SKILL_DISPLAY.items():
                if len(skill_key) <= 2:
                    # Short skills (R, Go, C#) need word boundaries
                    if re.search(r'\b' + re.escape(skill_key) + r'\b', bullet_lower):
                        _add_skill(display)
                elif skill_key in bullet_lower:
                    _add_skill(display)

    # 3) Scan summary/bio for known skills
    for key in [k for k in sections if any(w in k for w in ['summary', 'profile', 'objective'])]:
        summary_text = " ".join(sections[key]).lower()
        for skill_key, display in SKILL_DISPLAY.items():
            if len(skill_key) <= 2:
                if re.search(r'\b' + re.escape(skill_key) + r'\b', summary_text):
                    _add_skill(display)
            elif skill_key in summary_text:
                _add_skill(display)

    # 4) Scan header section (often has a tagline with skills)
    if 'header' in sections:
        header_text = " ".join(sections['header']).lower()
        for skill_key, display in SKILL_DISPLAY.items():
            if len(skill_key) <= 2:
                if re.search(r'\b' + re.escape(skill_key) + r'\b', header_text):
                    _add_skill(display)
            elif skill_key in header_text:
                _add_skill(display)

    result["skills"] = result["skills"][:30]

    # --- Parse Certifications ---
    cert_keys = [k for k in sections if 'cert' in k or 'license' in k or 'honor' in k or 'award' in k]
    for key in cert_keys:
        for cline in sections[key]:
            clean = re.sub(r'^[●•·▪◦○\-]\s*', '', cline).strip()
            if clean and len(clean) > 3:
                result["certifications"].append(clean)
    result["certifications"] = result["certifications"][:15]

    # --- Extract bio from summary/profile section ---
    bio_keys = [k for k in sections if any(w in k for w in ['summary', 'profile', 'objective'])]
    for key in bio_keys:
        bio_text = " ".join(sections[key]).strip()
        if bio_text and len(bio_text) > 20:
            result["bio"] = bio_text[:500]
            break

    logger.info(f"Basic resume parsing: {len(result['work_history'])} positions, {len(result['education'])} education, {len(result['skills'])} skills")
    return result


@router.post("/upload/resume")
async def upload_resume(
    file: UploadFile = File(...),
    current_user: dict = Depends(get_current_user)
):
    """Parse a PDF resume and return extracted profile data for autofill."""
    if current_user["role"] != "seeker":
        raise HTTPException(status_code=403, detail="Only job seekers can upload resumes")

    if file.content_type not in ALLOWED_RESUME_TYPES:
        raise HTTPException(status_code=400, detail="Only PDF and Word documents (.pdf, .doc, .docx) are accepted")

    contents = await file.read()
    if len(contents) > MAX_RESUME_SIZE:
        raise HTTPException(status_code=400, detail="Resume must be less than 10MB")

    # Extract text based on file type
    text = ""
    if file.content_type == "application/pdf":
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(io.BytesIO(contents))
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        except Exception as e:
            logger.error(f"PDF parsing failed: {e}")
            raise HTTPException(status_code=400, detail="Could not read PDF. Please ensure it is a valid, text-based PDF.")
    elif file.content_type in (
        "application/msword",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ):
        try:
            from docx import Document
            doc = Document(io.BytesIO(contents))
            for para in doc.paragraphs:
                if para.text.strip():
                    text += para.text + "\n"
            # Also extract text from tables (common in resumes)
            for table in doc.tables:
                for row in table.rows:
                    row_text = "\t".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text += row_text + "\n"
        except Exception as e:
            logger.error(f"Word document parsing failed: {e}")
            raise HTTPException(status_code=400, detail="Could not read Word document. Please ensure it is a valid .docx file.")

    if not text.strip():
        raise HTTPException(
            status_code=400,
            detail="Could not extract text from this file. It may be a scanned image. Please try a text-based document."
        )

    # Parse the extracted text using AI
    parsed = await _parse_resume_with_ai(text)

    logger.info(f"Resume parsed for user {current_user['id']}: {len(parsed.get('work_history', []))} positions, {len(parsed.get('education', []))} education, {len(parsed.get('skills', []))} skills, {len(parsed.get('certifications', []))} certs")
    return {"parsed": parsed, "raw_text_length": len(text)}
